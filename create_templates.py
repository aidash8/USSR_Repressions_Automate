import pandas as pd
import docx

def format_string(document: docx.Document, s: str):
    heading = document.add_heading('', level=1)
    run = heading.add_run(s)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(19)

def create_template(row, document):

    p = document.add_paragraph()
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = 0

    # Add name
    format_string(document=document, s=row['name'])
    # Add Age
    format_string(document=document, s=row['age'])
    # Add birth info
    format_string(document, row['birth_info'])
    # Add profession
    format_string(document, row['профессия / место работы'])
    # Add demise info
    format_string(document, row['demise_info']) 

    document.add_page_break()


def create_docs(name: str, df: pd.DataFrame):
    
    doc = docx.Document()

    for _,row in df.iterrows():
        create_template(row, doc)

    doc.save(f"{name}.docx")